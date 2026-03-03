"""
CAM (Credit Appraisal Memo) Generator
Populates a DOCX template with application data and exports a PDF.
Inserts a SHAP waterfall chart as an image.
"""
import io
import logging
import os
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Optional

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed — CAM generation disabled")


# Risk grade colours
GRADE_COLORS = {"A": "#27ae60", "B": "#2ecc71", "C": "#f39c12", "D": "#e67e22", "E": "#e74c3c"}


def _generate_shap_chart(shap_values: dict, output_path: str):
    """Create a horizontal SHAP waterfall bar chart and save to file."""
    if not shap_values:
        return

    names = list(shap_values.keys())
    values = [shap_values[n] for n in names]

    fig, ax = plt.subplots(figsize=(8, 4))
    colors = ["#27ae60" if v >= 0 else "#e74c3c" for v in values]
    bars = ax.barh(names, values, color=colors, edgecolor="white", height=0.6)

    for bar, val in zip(bars, values):
        ax.text(
            val + (max(abs(v) for v in values) * 0.02 if val >= 0 else -max(abs(v) for v in values) * 0.02),
            bar.get_y() + bar.get_height() / 2,
            f"{val:+.1f}",
            va="center",
            ha="left" if val >= 0 else "right",
            fontsize=9,
            color="#333333",
        )

    ax.axvline(0, color="#aaaaaa", linewidth=0.8)
    ax.set_xlabel("SHAP Impact on Credit Score", fontsize=10)
    ax.set_title("Feature Contributions (SHAP)", fontsize=11, fontweight="bold")
    ax.tick_params(axis="y", labelsize=9)
    fig.patch.set_facecolor("#f8f9fa")
    ax.set_facecolor("#f8f9fa")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches="tight")
    plt.close(fig)


def _add_section_heading(doc, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x1a, 0x73, 0xe8)


def _add_table_row(table, key: str, value: str):
    row = table.add_row()
    row.cells[0].text = key
    row.cells[1].text = str(value)
    row.cells[0].paragraphs[0].runs[0].bold = True


def generate_cam(
    application_id: str,
    company_name: str,
    gst_data: dict,
    bank_data: dict,
    itr_data: dict,
    reconciliation_data: dict,
    research_data: dict,
    scoring_result: dict,
    loan_decision: dict,
    output_dir: str = tempfile.gettempdir(),
) -> str:
    """
    Generate CAM as DOCX and attempt PDF export.
    Returns path to the generated file (PDF if LibreOffice available, else DOCX).
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed")

    doc = Document()

    # --- Title ---
    title = doc.add_heading(f"Credit Appraisal Memo", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(f"Application ID: {application_id}   |   Date: {datetime.now().strftime('%d %b %Y')}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x55, 0x5a, 0x64)

    doc.add_paragraph()

    # --- 1. Executive Summary ---
    _add_section_heading(doc, "1. Executive Summary")
    grade = scoring_result.get("risk_grade", "C")
    score = scoring_result.get("credit_score", 0)
    loan_amt = loan_decision.get("recommended_loan_amount", 0)
    rate = loan_decision.get("final_interest_rate", 10)
    approved = loan_decision.get("approved", False)

    summary_text = (
        f"This Credit Appraisal Memo evaluates the creditworthiness of {company_name}. "
        f"The applicant has been assigned a Credit Score of {score:.0f}/1000 with Risk Grade {grade}. "
        f"{'The application is APPROVED' if approved else 'The application is DECLINED'} "
        f"{'with a recommended loan limit of ₹' + f'{loan_amt:,.0f}' + ' at ' + str(rate) + '% per annum.' if approved else '.'}"
    )
    doc.add_paragraph(summary_text)

    # --- 2. Five Cs of Credit ---
    _add_section_heading(doc, "2. The Five Cs of Credit")

    ratios = scoring_result.get("financial_ratios", {})
    five_cs = [
        ("Character",
         f"Company: {company_name}. Research risk score: {research_data.get('overall_research_risk', 0):.2f}/1.0. "
         f"Court cases: {len(research_data.get('court_cases', []))}. "
         f"MCA status: {research_data.get('mca', {}).get('status', 'N/A')}."),
        ("Capacity",
         f"DSCR: {ratios.get('dscr', 0):.2f}. Interest Coverage: {ratios.get('interest_coverage', 0):.2f}. "
         f"Avg monthly bank credits: ₹{bank_data.get('average_monthly_credit', 0):,.0f}."),
        ("Capital",
         f"Equity: ₹{ratios.get('equity', 0):,.0f}. D/E ratio: {ratios.get('de_ratio', 0):.2f}. "
         f"Net profit (ITR): ₹{itr_data.get('net_profit', 0):,.0f}."),
        ("Collateral",
         f"Collateral value: ₹{loan_decision.get('collateral_value', 0):,.0f}. "
         f"Type: {loan_decision.get('collateral_type', 'N/A')}. "
         f"LTV: {loan_decision.get('ltv_ratio', 0)*100:.0f}%. "
         f"LTV Limit: ₹{loan_decision.get('collateral_ltv_limit', 0):,.0f}."),
        ("Conditions",
         f"GST-Bank reconciliation overall mismatch: "
         f"{reconciliation_data.get('overall_mismatch_ratio', 0):.1%}. "
         f"Anomalous months: {len(reconciliation_data.get('anomaly_months', []))}. "
         f"Reconciliation risk flag: {reconciliation_data.get('risk_flag', False)}."),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Parameter"
    hdr[1].text = "Assessment"
    for h in hdr:
        h.paragraphs[0].runs[0].bold = True

    for label, value in five_cs:
        _add_table_row(table, label, value)

    doc.add_paragraph()

    # --- 3. Financial Analysis ---
    _add_section_heading(doc, "3. Financial Analysis")

    _add_section_heading(doc, "3.1 Key Financial Ratios", level=2)
    ratio_table = doc.add_table(rows=1, cols=2)
    ratio_table.style = "Table Grid"
    ratio_table.rows[0].cells[0].text = "Ratio"
    ratio_table.rows[0].cells[1].text = "Value"
    for c in ratio_table.rows[0].cells:
        c.paragraphs[0].runs[0].bold = True

    ratio_display = [
        ("DSCR", f"{ratios.get('dscr', 0):.2f}"),
        ("Debt/Equity Ratio", f"{ratios.get('de_ratio', 0):.2f}"),
        ("Current Ratio", f"{ratios.get('current_ratio', 0):.2f}"),
        ("Interest Coverage", f"{ratios.get('interest_coverage', 0):.2f}"),
        ("Gross Margin", f"{ratios.get('gross_margin', 0):.1%}"),
        ("Credit Score", f"{score:.0f} / 1000"),
        ("Risk Grade", grade),
    ]
    for label, val in ratio_display:
        _add_table_row(ratio_table, label, val)

    doc.add_paragraph()

    _add_section_heading(doc, "3.2 GST-Bank Reconciliation", level=2)
    doc.add_paragraph(reconciliation_data.get("summary", "N/A"))

    _add_section_heading(doc, "3.3 External Research", level=2)
    doc.add_paragraph(research_data.get("news_summary", "") + "\n" + research_data.get("legal_summary", ""))
    if research_data.get("mca_flags"):
        for flag in research_data["mca_flags"]:
            doc.add_paragraph(f"• {flag}", style="List Bullet")

    # --- 4. SHAP Analysis ---
    _add_section_heading(doc, "4. Credit Score Explanation (SHAP)")
    shap_vals = scoring_result.get("shap_values", {})
    if shap_vals:
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tf:
            shap_img_path = tf.name
        try:
            _generate_shap_chart(shap_vals, shap_img_path)
            doc.add_picture(shap_img_path, width=Inches(6))
        except Exception as e:
            logger.warning(f"SHAP chart insert failed: {e}")
            doc.add_paragraph("(SHAP chart unavailable)")
        finally:
            if os.path.exists(shap_img_path):
                os.unlink(shap_img_path)
    else:
        doc.add_paragraph("No SHAP values available.")

    # Rule flags
    if scoring_result.get("rule_flags"):
        doc.add_paragraph()
        doc.add_paragraph("Rule-Based Flags:")
        for flag in scoring_result["rule_flags"]:
            doc.add_paragraph(f"• {flag}", style="List Bullet")

    # --- 5. Loan Recommendation ---
    _add_section_heading(doc, "5. Loan Recommendation")
    rec_table = doc.add_table(rows=1, cols=2)
    rec_table.style = "Table Grid"
    rec_table.rows[0].cells[0].text = "Parameter"
    rec_table.rows[0].cells[1].text = "Value"
    for c in rec_table.rows[0].cells:
        c.paragraphs[0].runs[0].bold = True

    rec_rows = [
        ("Decision", "APPROVED ✓" if approved else "DECLINED ✗"),
        ("Recommended Loan Amount", f"₹{loan_amt:,.0f}"),
        ("Interest Rate", f"{rate}% per annum"),
        ("Cash-Flow Capacity", f"₹{loan_decision.get('cash_flow_capacity', 0):,.0f}"),
        ("Collateral LTV Limit", f"₹{loan_decision.get('collateral_ltv_limit', 0):,.0f}"),
    ]
    for label, val in rec_rows:
        _add_table_row(rec_table, label, val)

    doc.add_paragraph()
    if loan_decision.get("reasons"):
        doc.add_paragraph("Rationale:")
        for reason in loan_decision["reasons"]:
            doc.add_paragraph(f"• {reason}", style="List Bullet")

    # --- Footer ---
    doc.add_paragraph()
    footer = doc.add_paragraph(
        "This memo is computer-generated by the CreditSight AI Engine. "
        "Decisions are subject to credit committee review and applicable regulatory norms."
    )
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # --- Save DOCX ---
    docx_path = os.path.join(output_dir, f"cam_{application_id}.docx")
    doc.save(docx_path)
    logger.info(f"CAM DOCX saved: {docx_path}")

    # --- Try PDF conversion ---
    pdf_path = docx_path.replace(".docx", ".pdf")
    try:
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", output_dir, docx_path],
            capture_output=True, timeout=60
        )
        if result.returncode == 0 and os.path.exists(pdf_path):
            logger.info(f"CAM PDF created: {pdf_path}")
            return pdf_path
    except Exception as e:
        logger.warning(f"LibreOffice PDF conversion failed ({e}), trying docx2pdf")

    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        if os.path.exists(pdf_path):
            return pdf_path
    except Exception as e:
        logger.warning(f"docx2pdf failed ({e}) — returning DOCX")

    return docx_path
